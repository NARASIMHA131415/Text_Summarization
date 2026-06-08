# =============================================
# file_extractor.py
# Extracts text from PDF, DOC, DOCX, XLS, XLSX, TXT
# =============================================

import os
import fitz                          # PyMuPDF  → PDF
import docx                          # python-docx → DOC/DOCX
import openpyxl                      # openpyxl  → XLSX
import xlrd                          # xlrd      → XLS


# =============================================
# MAIN EXTRACTOR - routes to correct extractor
# =============================================
def extract_text(file_path: str) -> str:
    """
    Main function to extract text from any supported file.

    Args:
        file_path (str): Full path to the uploaded file

    Returns:
        str: Extracted plain text from the file
    """

    # Get file extension
    extension = os.path.splitext(file_path)[1].lower()

    # Route to correct extractor
    extractors = {
        '.pdf':  extract_from_pdf,
        '.doc':  extract_from_doc,
        '.docx': extract_from_docx,
        '.xls':  extract_from_xls,
        '.xlsx': extract_from_xlsx,
        '.txt':  extract_from_txt,
    }

    if extension not in extractors:
        raise ValueError(f"Unsupported file type: {extension}")

    # Call the correct extractor
    extracted_text = extractors[extension](file_path)

    # Clean and return
    return clean_text(extracted_text)


# =============================================
# PDF EXTRACTOR  (PyMuPDF)
# =============================================
def extract_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file using PyMuPDF.

    Args:
        file_path (str): Path to the PDF file

    Returns:
        str: Extracted text from all pages
    """
    text = []

    try:
        # Open PDF document
        pdf_document = fitz.open(file_path)

        # Loop through all pages
        for page_number in range(len(pdf_document)):
            page        = pdf_document.load_page(page_number)
            page_text   = page.get_text("text")   # Extract plain text
            text.append(page_text)

        pdf_document.close()

    except Exception as e:
        raise RuntimeError(f"Failed to extract text from PDF: {str(e)}")

    return "\n".join(text)


# =============================================
# DOC EXTRACTOR  (python-docx handles DOCX)
# For old .doc we convert via python-docx
# =============================================
def extract_from_doc(file_path: str) -> str:
    """
    Extract text from a .doc file.
    Note: python-docx works best with .docx
          For old .doc, we attempt with python-docx.

    Args:
        file_path (str): Path to the DOC file

    Returns:
        str: Extracted text
    """
    try:
        return extract_from_docx(file_path)
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from DOC: {str(e)}")


# =============================================
# DOCX EXTRACTOR  (python-docx)
# =============================================
def extract_from_docx(file_path: str) -> str:
    """
    Extract text from a .docx file using python-docx.

    Args:
        file_path (str): Path to the DOCX file

    Returns:
        str: Extracted text from all paragraphs
    """
    text = []

    try:
        doc = docx.Document(file_path)

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():          # Skip empty paragraphs
                text.append(paragraph.text)

        # Extract text from tables too
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)

    except Exception as e:
        raise RuntimeError(f"Failed to extract text from DOCX: {str(e)}")

    return "\n".join(text)


# =============================================
# XLSX EXTRACTOR  (openpyxl)
# =============================================
def extract_from_xlsx(file_path: str) -> str:
    """
    Extract text from a .xlsx Excel file using openpyxl.

    Args:
        file_path (str): Path to the XLSX file

    Returns:
        str: Extracted text from all sheets and cells
    """
    text = []

    try:
        workbook = openpyxl.load_workbook(file_path, read_only=True, data_only=True)

        # Loop through all sheets
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text.append(f"[Sheet: {sheet_name}]")

            # Loop through all rows and cells
            for row in sheet.iter_rows(values_only=True):
                row_text = []
                for cell in row:
                    if cell is not None and str(cell).strip():
                        row_text.append(str(cell))

                if row_text:
                    text.append(" | ".join(row_text))

        workbook.close()

    except Exception as e:
        raise RuntimeError(f"Failed to extract text from XLSX: {str(e)}")

    return "\n".join(text)


# =============================================
# XLS EXTRACTOR  (xlrd for old Excel format)
# =============================================
def extract_from_xls(file_path: str) -> str:
    """
    Extract text from a .xls Excel file using xlrd.

    Args:
        file_path (str): Path to the XLS file

    Returns:
        str: Extracted text from all sheets and cells
    """
    text = []

    try:
        workbook = xlrd.open_workbook(file_path)

        # Loop through all sheets
        for sheet_index in range(workbook.nsheets):
            sheet = workbook.sheet_by_index(sheet_index)
            text.append(f"[Sheet: {sheet.name}]")

            # Loop through all rows
            for row_index in range(sheet.nrows):
                row_values = sheet.row_values(row_index)
                row_text   = [str(cell).strip() for cell in row_values if str(cell).strip()]

                if row_text:
                    text.append(" | ".join(row_text))

    except Exception as e:
        raise RuntimeError(f"Failed to extract text from XLS: {str(e)}")

    return "\n".join(text)


# =============================================
# TXT EXTRACTOR  (built-in)
# =============================================
def extract_from_txt(file_path: str) -> str:
    """
    Extract text from a plain .txt file.

    Args:
        file_path (str): Path to the TXT file

    Returns:
        str: Plain text content
    """
    try:
        # Try UTF-8 first, fallback to latin-1
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()

    except Exception as e:
        raise RuntimeError(f"Failed to extract text from TXT: {str(e)}")


# =============================================
# TEXT CLEANER
# =============================================
def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text.

    Args:
        text (str): Raw extracted text

    Returns:
        str: Cleaned text
    """
    import re

    if not text:
        return ""

    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)     # Max 2 newlines
    text = re.sub(r' {2,}',  ' ',    text)     # Max 1 space
    text = re.sub(r'\t',     ' ',    text)     # Replace tabs

    # Remove non-printable characters
    text = re.sub(r'[^\x20-\x7E\n]', ' ', text)

    return text.strip()


# =============================================
# UTILITY - Get file info
# =============================================
def get_file_info(file_path: str) -> dict:
    """
    Get basic info about the uploaded file.

    Args:
        file_path (str): Path to the file

    Returns:
        dict: File info (name, size, extension)
    """
    file_name  = os.path.basename(file_path)
    file_size  = os.path.getsize(file_path)
    extension  = os.path.splitext(file_path)[1].lower()

    return {
        "file_name"  : file_name,
        "file_size"  : f"{round(file_size / 1024, 2)} KB",
        "file_type"  : extension.replace('.', '').upper()
    }       